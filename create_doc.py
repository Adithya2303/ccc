from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Skill Exchange Project Documentation', 0)

def add_section(title, content):
    doc.add_heading(title, level=1)
    if isinstance(content, list):
        for item in content:
            doc.add_paragraph(item, style='List Number')
    else:
        doc.add_paragraph(content)

def add_subsection(title, content):
    doc.add_heading(title, level=2)
    if isinstance(content, list):
        for item in content:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph(content)

# 1. Functional Requirements
add_section('1. Functional Requirements', [
    'User Registration and Login (JWT-based authentication)',
    'Profile Management (Add/Edit skills and skills to learn)',
    'Skill Matching (Find users who can teach what you want to learn and vice versa)',
    'Real-time Chat between matched users',
    'Session Management (Token validation on app load, auto-logout if invalid)',
])

# 2. Database Design
add_section('2. Database Design', '')
add_subsection('User Collection', [
    'username: String, required, unique',
    'email: String, required, unique',
    'password: String, required (hashed)',
    'skills: [String]',
    'wantsToLearn: [String]',
])
add_subsection('Chat Collection', [
    'users: [ObjectId] (references User)',
    'messages: [ { sender: ObjectId (User), content: String, timestamp: Date } ]',
])

# 3. Project Structure & REST API Design
add_section('3. Project Structure & REST API Design', '')
add_subsection('Backend Structure', [
    'controllers/: Business logic for authentication, chat, skills, and matching',
    'routes/: Express route definitions for each feature',
    'models/: Mongoose schemas for User and Chat',
    'middleware/: JWT authentication middleware',
    'index.js: Main server file, sets up Express, Socket.io, and MongoDB connection',
])
add_subsection('Frontend Structure', [
    'src/: React components (App.jsx, Login.jsx, Register.jsx, Skills.jsx)',
    'public/: Static assets and images',
])
add_subsection('Main REST API Endpoints', [
    'POST /api/auth/register - Register a new user',
    'POST /api/auth/login - Login and receive JWT token',
    'GET /api/auth/check - Validate JWT token',
    'GET /api/skills - Get user skills and wants',
    'POST /api/skills - Update user skills and wants',
    'GET /api/match - Find matching users',
    'GET /api/chats/:otherUserId - Get chat history with another user',
    'POST /api/chats/:otherUserId - Send a message to another user',
])

# 4. UI Design
add_section('4. UI Design', [
    'Login/Register screen: Allows user to register or login',
    'Profile screen: Shows and allows editing of skills and skills to learn',
    'Match screen: Shows list of matching users and allows starting a chat',
    'Chat window: Real-time chat interface with matched users',
    'Logout button: Allows user to log out and clear session',
])

# Save the document
doc.save('SkillExchange_Documentation.docx') 